"""
Document generation & extraction router.
Handles .docx, .xlsx, .pdf creation and text extraction from uploaded files.
"""

from __future__ import annotations

import io
import json
import os
import re
import uuid
from datetime import datetime
from pathlib import Path

from fastapi import APIRouter, Depends, HTTPException, UploadFile, File as FastAPIFile
from fastapi.responses import Response
from pydantic import BaseModel
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from backend.database import get_db, File as FileModel, Project

router = APIRouter(prefix="/api/generate", tags=["documents"])

UPLOAD_DIR = Path(__file__).parent.parent / "uploads"


# ---------------------------------------------------------------------------
# Request schemas
# ---------------------------------------------------------------------------

class DocxRequest(BaseModel):
    title: str = "Document"
    content: str              # Markdown text content


class XlsxRequest(BaseModel):
    title: str = "Sheet"
    sheets: list[dict]        # [{"name": "Sheet1", "headers": [...], "rows": [[...], ...]}, ...]


class PdfRequest(BaseModel):
    title: str = "Document"
    content: str              # Markdown text content
    author: str = "ProjektLLM"


class ExtractRequest(BaseModel):
    file_id: str


# ---------------------------------------------------------------------------
# DocX generation
# ---------------------------------------------------------------------------

def _build_docx(title: str, content: str) -> io.BytesIO:
    """Convert markdown-ish text to .docx and return as BytesIO."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse simple markdown: headings, bold, italic, lists, paragraphs
    lines = content.split("\n")
    for line in lines:
        stripped = line.strip()

        if not stripped:
            continue

        # Headings
        h_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if h_match:
            level = len(h_match.group(1))
            doc.add_heading(h_match.group(2), level=level)
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}$", stripped):
            doc.add_paragraph("_" * 60)
            continue

        # Unordered list item
        if re.match(r"^[\s]*[-*+]\s+", stripped):
            text = re.sub(r"^[\s]*[-*+]\s+", "", stripped)
            p = doc.add_paragraph(text, style="List Bullet")
            _apply_inline_formatting(p)
            continue

        # Numbered list item
        if re.match(r"^\s*\d+[.)]\s+", stripped):
            text = re.sub(r"^\s*\d+[.)]\s+", "", stripped)
            p = doc.add_paragraph(text, style="List Number")
            _apply_inline_formatting(p)
            continue

        # Regular paragraph — preserve line breaks within
        p = doc.add_paragraph(stripped)
        _apply_inline_formatting(p)

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _apply_inline_formatting(paragraph):
    """Apply basic **bold** and *italic* within a paragraph."""
    import docx
    text = paragraph.text
    if "**" not in text and "*" not in text:
        return
    # Clear and rebuild runs
    paragraph.clear()
    parts = re.split(r"(\*\*\*.*?\*\*\*|\*\*.*?\*\*|\*.*?\*)", text)
    for part in parts:
        if part.startswith("***") and part.endswith("***"):
            run = paragraph.add_run(part[3:-3])
            run.bold = True
            run.italic = True
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


# ---------------------------------------------------------------------------
# XLSX generation
# ---------------------------------------------------------------------------

def _build_xlsx(request: XlsxRequest) -> io.BytesIO:
    """Generate .xlsx from structured sheet data."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    # Remove default sheet
    wb.remove(wb.active)

    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF", size=11)
    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    for sheet_def in request.sheets:
        name = sheet_def.get("name", "Sheet")
        headers = sheet_def.get("headers", [])
        rows = sheet_def.get("rows", [])

        ws = wb.create_sheet(title=name[:31])  # Excel limit

        # Write headers
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.border = thin_border

        # Write data rows
        for row_idx, row_data in enumerate(rows, 2):
            for col_idx, value in enumerate(row_data, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.border = thin_border
                cell.alignment = Alignment(vertical="center")

        # Auto-adjust column widths
        for col_idx, header in enumerate(headers, 1):
            col_letter = get_column_letter(col_idx)
            max_len = len(str(header))
            for row_data in rows:
                if col_idx <= len(row_data):
                    max_len = max(max_len, len(str(row_data[col_idx - 1])))
            ws.column_dimensions[col_letter].width = min(max_len + 3, 60)

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# PDF generation
# ---------------------------------------------------------------------------

def _build_pdf(title: str, content: str, author: str = "ProjektLLM") -> io.BytesIO:
    """Generate PDF from markdown-ish text using ReportLab."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm, mm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, ListFlowable, ListItem,
        HRFlowable, PageBreak,
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY

    buf = io.BytesIO()

    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        title=title,
        author=author,
    )

    styles = getSampleStyleSheet()
    # Custom styles
    styles.add(ParagraphStyle(
        "Title2", parent=styles["Title"], fontSize=24, spaceAfter=20,
        alignment=TA_CENTER, textColor=HexColor("#1a1a2e"),
    ))
    styles.add(ParagraphStyle(
        "Heading1Custom", parent=styles["Heading1"], fontSize=18,
        spaceBefore=16, spaceAfter=8, textColor=HexColor("#16213e"),
    ))
    styles.add(ParagraphStyle(
        "Heading2Custom", parent=styles["Heading2"], fontSize=14,
        spaceBefore=12, spaceAfter=6, textColor=HexColor("#0f3460"),
    ))
    styles.add(ParagraphStyle(
        "BodyCustom", parent=styles["Normal"], fontSize=10.5,
        leading=15, alignment=TA_JUSTIFY, spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        "BulletCustom", parent=styles["Normal"], fontSize=10.5,
        leading=15, leftIndent=20, bulletIndent=10, spaceAfter=4,
    ))
    styles.add(ParagraphStyle(
        "CodeBlock", parent=styles["Normal"], fontSize=8.5,
        leading=11, leftIndent=15, fontName="Courier",
        backColor=HexColor("#f5f5f5"), spaceAfter=8,
    ))

    story = []
    lines = content.split("\n")
    in_code_block = False
    code_buf_lines = []

    for line in lines:
        stripped = line.strip()

        # Code block toggle
        if stripped.startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_buf_lines)
                story.append(Paragraph(
                    code_text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br/>"),
                    styles["CodeBlock"],
                ))
                code_buf_lines = []
                in_code_block = False
            else:
                in_code_block = True
            continue

        if in_code_block:
            code_buf_lines.append(line)
            continue

        if not stripped:
            story.append(Spacer(1, 6))
            continue

        # Heading
        h_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if h_match:
            level = len(h_match.group(1))
            text = h_match.group(2)
            if level == 1:
                story.append(Paragraph(_escape(text), styles["Heading1Custom"]))
            elif level == 2:
                story.append(Paragraph(_escape(text), styles["Heading2Custom"]))
            else:
                story.append(Paragraph(_escape(text), styles["Heading3"]))
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}$", stripped):
            story.append(HRFlowable(width="100%", thickness=1, color=HexColor("#cccccc")))
            continue

        # Bullet list
        if re.match(r"^[\s]*[-*+]\s+", stripped):
            text = re.sub(r"^[\s]*[-*+]\s+", "", stripped)
            bullet = _inline_markdown_to_html(text)
            story.append(ListFlowable(
                [ListItem(Paragraph(bullet, styles["BulletCustom"]), bulletColor=HexColor("#4472C4"))],
                bulletType="bullet",
                start="•",
                leftIndent=20,
            ))
            continue

        # Numbered list
        num_match = re.match(r"^\s*(\d+)[.)]\s+(.+)", stripped)
        if num_match:
            text = num_match.group(2)
            item = _inline_markdown_to_html(text)
            story.append(ListFlowable(
                [ListItem(Paragraph(item, styles["BulletCustom"]))],
                bulletType="bullet",
                start=num_match.group(1),
                leftIndent=20,
            ))
            continue

        # Regular paragraph
        html = _inline_markdown_to_html(stripped)
        story.append(Paragraph(html, styles["BodyCustom"]))

    doc.build(story)
    buf.seek(0)
    return buf


def _escape(text: str) -> str:
    """Escape XML special chars for ReportLab."""
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _inline_markdown_to_html(text: str) -> str:
    """Convert inline markdown (**bold**, *italic*, `code`, [link](url)) to HTML."""
    text = _escape(text)
    # **bold** and ***bold italic***
    text = re.sub(r"\*\*\*(.*?)\*\*\*", r"<b><i>\1</i></b>", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"\*(.*?)\*", r"<i>\1</i>", text)
    # `code`
    text = re.sub(r"`(.*?)`", r"<font face='Courier'><b>\1</b></font>", text)
    # [text](url)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r'<a href="\2" color="#4472C4">\1</a>', text)
    return text


# ---------------------------------------------------------------------------
# Text extraction from uploaded files
# ---------------------------------------------------------------------------

async def extract_text_from_file(file_path: str) -> str:
    """Extract readable text from PDF, DOCX, or XLSX files."""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf_text(file_path)
    elif ext == ".docx":
        return _extract_docx_text(file_path)
    elif ext == ".xlsx":
        return _extract_xlsx_text(file_path)
    else:
        # Fallback: try as plain text
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                return f.read()
        except Exception:
            return f"[Cannot extract text from {ext} file]"


def _extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF using PyMuPDF."""
    import fitz
    doc = fitz.open(file_path)
    pages = []
    for page_num, page in enumerate(doc):
        text = page.get_text()
        if text.strip():
            pages.append(f"--- Page {page_num + 1} ---\n{text}")
    doc.close()
    return "\n\n".join(pages) if pages else "[No extractable text found in PDF]"


def _extract_docx_text(file_path: str) -> str:
    """Extract text from DOCX."""
    from docx import Document
    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        parts.append("\n".join(rows))
    return "\n\n".join(parts) if parts else "[No text content found in DOCX]"


def _extract_xlsx_text(file_path: str) -> str:
    """Extract text from XLSX."""
    from openpyxl import load_workbook
    wb = load_workbook(file_path, read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_text = []
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = "\t".join(cells)
            if line.strip():
                rows_text.append(line)
        if rows_text:
            parts.append(f"=== Sheet: {sheet_name} ===\n" + "\n".join(rows_text))
    wb.close()
    return "\n\n".join(parts) if parts else "[No data found in XLSX]"


# ---------------------------------------------------------------------------
# API endpoints
# ---------------------------------------------------------------------------

@router.post("/docx")
async def generate_docx(body: DocxRequest):
    """Generate a .docx file from markdown content."""
    buf = _build_docx(body.title, body.content)
    filename = f"{body.title.replace(' ', '_')}.docx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/xlsx")
async def generate_xlsx(body: XlsxRequest):
    """Generate a .xlsx file from structured sheet data."""
    buf = _build_xlsx(body)
    filename = f"{body.title.replace(' ', '_')}.xlsx"
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/pdf")
async def generate_pdf(body: PdfRequest):
    """Generate a PDF from markdown content."""
    buf = _build_pdf(body.title, body.content, body.author)
    filename = f"{body.title.replace(' ', '_')}.pdf"
    return Response(
        content=buf.read(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.post("/extract")
async def extract_file_content(
    body: ExtractRequest,
    db: AsyncSession = Depends(get_db),
):
    """Extract text content from an uploaded file (PDF, DOCX, XLSX, or text)."""
    file_record = await db.get(FileModel, body.file_id)
    if not file_record or not file_record.filepath:
        raise HTTPException(404, "File not found")

    path = file_record.filepath
    if not os.path.exists(path):
        raise HTTPException(404, "File data not found on disk")

    text = await extract_text_from_file(path)
    return {
        "file_id": body.file_id,
        "name": file_record.name,
        "kind": file_record.kind,
        "text": text,
        "length": len(text),
    }


@router.post("/extract-upload")
async def extract_uploaded_file(
    file: UploadFile = FastAPIFile(...),
):
    """Extract text from a directly uploaded file (without saving to DB)."""
    content = await file.read()
    ext = Path(file.filename or "upload.txt").suffix.lower()

    # Save to temp to process
    tmp = UPLOAD_DIR / f"_extract_{uuid.uuid4().hex}{ext}"
    tmp.parent.mkdir(parents=True, exist_ok=True)
    tmp.write_bytes(content)

    try:
        text = await extract_text_from_file(str(tmp))
    finally:
        if tmp.exists():
            tmp.unlink()

    return {
        "name": file.filename or "unnamed",
        "text": text,
        "length": len(text),
    }
