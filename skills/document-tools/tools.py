#!/usr/bin/env python3
"""
Document Tools — generate and manipulate DOCX, XLSX, PDF, Markdown.

Usage:
  python tools.py docx create <output.docx> [--title T] [--heading H] [--body B] [--json data.json]
  python tools.py docx edit <file.docx> [--replace old=new ...]
  python tools.py docx to-md <file.docx> [--output README.md]
  python tools.py xlsx create <output.xlsx> [--csv data.csv] [--json data.json]
  python tools.py xlsx edit <file.xlsx> [--set "Sheet!A1=value"]
  python tools.py xlsx convert <input> [--output out.xlsx]
  python tools.py pdf read <file.pdf> [--output text.txt]
  python tools.py pdf merge --inputs a.pdf b.pdf -o out.pdf
  python tools.py pdf split <file.pdf> [--pages 1-3,5]
  python tools.py pdf watermark <file.pdf> -w "DRAFT" -o out.pdf
  python tools.py pdf encrypt <file.pdf> -p password -o out.pdf
  python tools.py md to-docx <file.md> [--styled] -o out.docx
"""
import argparse, csv, io, json, os, re, sys, tempfile, uuid
from pathlib import Path
from datetime import datetime
from xml.etree import ElementTree as ET

# ── DOCX ───────────────────────────────────────────────────────────────────────
def docx_create(args):
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml

    doc = Document()

    # Style config
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)

    if args.title:
        p = doc.add_heading(args.title, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if args.json:
        with open(args.json) as f:
            data = json.load(f)
        for item in data:
            t = item.get("type", "paragraph")
            if t == "heading":
                level = item.get("level", 1)
                doc.add_heading(item["text"], level=level)
            elif t == "paragraph":
                p = doc.add_paragraph(item.get("text", ""))
                if item.get("bold"):
                    for run in p.runs:
                        run.bold = True
                if item.get("alignment") == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif t == "table":
                headers = item.get("headers", [])
                rows = item.get("rows", [])
                table = doc.add_table(rows=1 + len(rows), cols=len(headers))
                table.style = 'Light Grid Accent 1'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for i, h in enumerate(headers):
                    cell = table.rows[0].cells[i]
                    cell.text = h
                    cell.paragraphs[0].runs[0].bold = True if cell.paragraphs[0].runs else True
                for ri, row in enumerate(rows):
                    for ci, val in enumerate(row):
                        table.rows[ri + 1].cells[ci].text = str(val)
            elif t == "image":
                path = item.get("path", "")
                if os.path.exists(path):
                    width = Inches(item.get("width", 5))
                    doc.add_picture(path, width=width)
                    last = doc.paragraphs[-1]
                    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif t == "page_break":
                doc.add_page_break()
            elif t == "list":
                for li in item.get("items", []):
                    doc.add_paragraph(li, style='List Bullet')
            elif t == "numbered_list":
                for li in item.get("items", []):
                    doc.add_paragraph(li, style='List Number')

    if args.heading:
        doc.add_heading(args.heading, level=1)
    if args.body:
        doc.add_paragraph(args.body)

    # Footer with page numbers
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)

    doc.save(args.output)
    print(json.dumps({"status": "ok", "file": args.output}))


def docx_edit(args):
    from docx import Document
    doc = Document(args.input)

    if args.replace:
        for repl in args.replace:
            if "=" not in repl:
                continue
            old, new = repl.split("=", 1)
            for p in doc.paragraphs:
                for run in p.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            for run in p.runs:
                                if old in run.text:
                                    run.text = run.text.replace(old, new)

    if args.append:
        doc.add_paragraph(args.append)

    doc.save(args.input)
    print(json.dumps({"status": "ok", "file": args.input}))


def docx_to_md(args):
    from docx import Document
    doc = Document(args.input)
    lines = []
    for p in doc.paragraphs:
        if p.style.name.startswith("Heading"):
            level = p.style.name.split()[-1]
            lines.append(f"{'#' * int(level)} {p.text}")
        else:
            lines.append(p.text)

    # Tables
    for table in doc.tables:
        md_rows = []
        for ri, row in enumerate(table.rows):
            cells = [cell.text.replace("\n", " ") for cell in row.cells]
            md_rows.append("| " + " | ".join(cells) + " |")
            if ri == 0:
                md_rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
        lines.extend(["", *md_rows, ""])

    output = args.output or re.sub(r"\.docx$", ".md", args.input, flags=re.I)
    Path(output).write_text("\n".join(lines), encoding="utf-8")
    print(json.dumps({"status": "ok", "file": output, "length": len(lines)}))


# ── XLSX ───────────────────────────────────────────────────────────────────────
def xlsx_create(args):
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    ws = wb.active
    ws.title = "Sheet1"

    thin = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    has_data = False

    if args.json:
        with open(args.json) as f:
            data = json.load(f)
        if isinstance(data, dict):
            data = [data]
        if data:
            has_data = True
            headers = list(data[0].keys())
            ws.append(headers)
            for col in range(1, len(headers) + 1):
                cell = ws.cell(row=1, column=col)
                cell.font = Font(bold=True)
                cell.border = thin
                cell.alignment = Alignment(horizontal="center")
            for row in data:
                ws.append([row.get(h, "") for h in headers])

    if args.csv:
        with open(args.csv, newline="") as f:
            reader = csv.reader(f)
            for row in reader:
                ws.append(row)
        has_data = True

    if args.rows:
        for row in args.rows:
            ws.append(row)

    if not has_data:
        ws.append(["No", "Item", "Value"])
        ws.append([1, "Example", "data"])

    # Auto-width
    for col in range(1, ws.max_column + 1):
        max_len = 0
        for row in range(1, min(ws.max_row + 1, 50)):
            val = ws.cell(row=row, column=col).value
            if val:
                max_len = max(max_len, len(str(val)))
        ws.column_dimensions[get_column_letter(col)].width = min(max_len + 3, 50)

    wb.save(args.output)
    print(json.dumps({"status": "ok", "file": args.output, "rows": ws.max_row, "cols": ws.max_column}))


def xlsx_edit(args):
    from openpyxl import load_workbook
    wb = load_workbook(args.input)
    ws = wb.active

    if args.set:
        for s in args.set:
            if "=" not in s:
                continue
            cell_ref, value = s.split("=", 1)
            if "!" in cell_ref:
                sheet, cell = cell_ref.split("!", 1)
                wb[sheet][cell] = value
            else:
                ws[cell_ref] = value

    if args.add_row:
        parts = args.add_row.split(",")
        ws.append(parts)

    wb.save(args.input)
    print(json.dumps({"status": "ok", "file": args.input}))


def xlsx_convert(args):
    from openpyxl import Workbook, load_workbook

    ext = Path(args.input).suffix.lower()
    output = args.output or re.sub(r"\.[^.]+$", ".xlsx", args.input)

    if ext == ".csv":
        wb = Workbook()
        ws = wb.active
        with open(args.input, newline="", encoding="utf-8-sig") as f:
            reader = csv.reader(f)
            for row in reader:
                ws.append(row)
        wb.save(output)
    elif ext == ".tsv":
        wb = Workbook()
        ws = wb.active
        with open(args.input, newline="") as f:
            reader = csv.reader(f, delimiter="\t")
            for row in reader:
                ws.append(row)
        wb.save(output)
    elif ext == ".xlsx":
        wb = load_workbook(args.input)
        ws = wb.active
        output = re.sub(r"\.[^.]+$", ".csv", args.input)
        with open(output, "w", newline="") as f:
            writer = csv.writer(f)
            for row in ws.iter_rows(values_only=True):
                writer.writerow(row)
    else:
        print(json.dumps({"status": "error", "message": f"Unsupported format: {ext}"}))
        return

    print(json.dumps({"status": "ok", "file": output}))


# ── PDF ────────────────────────────────────────────────────────────────────────
def pdf_read(args):
    import pdfplumber
    result = {"pages": [], "tables": [], "text": ""}

    with pdfplumber.open(args.input) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            tables = page.extract_tables() or []
            result["pages"].append({"num": i + 1, "text": text, "table_count": len(tables)})
            for table in tables:
                result["tables"].append(table)
            result["text"] += f"\n--- Page {i + 1} ---\n{text}"

    result["page_count"] = len(result["pages"])
    result["table_count"] = len(result["tables"])

    if args.output:
        Path(args.output).write_text(result["text"], encoding="utf-8")
        result["saved_to"] = args.output

    print(json.dumps(result, ensure_ascii=False))


def pdf_merge(args):
    from pypdf import PdfWriter
    merger = PdfWriter()
    for path in args.inputs:
        merger.append(path)
    output = args.output or f"merged_{uuid.uuid4().hex[:8]}.pdf"
    merger.write(output)
    merger.close()
    print(json.dumps({"status": "ok", "file": output, "pages": len(merger.pages)}))


def pdf_split(args):
    from pypdf import PdfReader, PdfWriter
    reader = PdfReader(args.input)
    output = args.output or f"split_{uuid.uuid4().hex[:8]}.pdf"

    if args.pages:
        ranges = []
        for part in args.pages.split(","):
            if "-" in part:
                start, end = part.split("-")
                ranges.extend(range(int(start) - 1, int(end)))
            else:
                ranges.append(int(part) - 1)
    else:
        ranges = list(range(len(reader.pages)))

    writer = PdfWriter()
    for i in ranges:
        if i < len(reader.pages):
            writer.add_page(reader.pages[i])

    writer.write(output)
    writer.close()
    print(json.dumps({"status": "ok", "file": output, "pages": len(ranges)}))


def pdf_watermark(args):
    from pypdf import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas

    reader = PdfReader(args.input)
    writer = PdfWriter()

    # Create watermark overlay
    overlay_path = tempfile.mktemp(suffix=".pdf")
    c = canvas.Canvas(overlay_path, pagesize=(612, 792))
    c.saveState()
    c.setFont("Helvetica", 60)
    c.setFillColorRGB(0.8, 0.8, 0.8, 0.3)
    c.translate(306, 396)
    c.rotate(45)
    c.drawCentredString(0, 0, args.watermark_text)
    c.restoreState()
    c.save()

    watermark_reader = PdfReader(overlay_path)
    watermark_page = watermark_reader.pages[0]

    for page in reader.pages:
        page.merge_page(watermark_page)
        writer.add_page(page)

    output = args.output or f"watermarked_{uuid.uuid4().hex[:8]}.pdf"
    writer.write(output)
    writer.close()
    os.unlink(overlay_path)
    print(json.dumps({"status": "ok", "file": output}))


def pdf_encrypt(args):
    from pypdf import PdfReader, PdfWriter

    reader = PdfReader(args.input)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt(args.password)

    output = args.output or f"encrypted_{uuid.uuid4().hex[:8]}.pdf"
    writer.write(output)
    writer.close()
    print(json.dumps({"status": "ok", "file": output, "encrypted": True}))


# ── Markdown ───────────────────────────────────────────────────────────────────
def md_to_docx(args):
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    md_text = Path(args.input).read_text(encoding="utf-8")
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    lines = md_text.split("\n")
    i = 0
    in_table = False
    table_rows = []
    table_headers = []

    while i < len(lines):
        line = lines[i]

        # Headings
        hm = re.match(r"^(#{1,6})\s+(.+)$", line)
        if hm:
            level = len(hm.group(1))
            doc.add_heading(hm.group(2), level=level)
            i += 1
            continue

        # Table separator
        if re.match(r"^\|[\s\-:|]+\|$", line):
            i += 1
            continue

        # Table row
        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.strip("|").split("|")]
            if not in_table:
                table_headers = cells
                in_table = True
            else:
                table_rows.append(cells)
            i += 1
            if i >= len(lines) or not lines[i].startswith("|"):
                in_table = False
                if table_rows:
                    table = doc.add_table(rows=1 + len(table_rows), cols=len(table_headers))
                    table.style = 'Light Grid Accent 1'
                    for ci, h in enumerate(table_headers):
                        table.rows[0].cells[ci].text = h
                    for ri, row in enumerate(table_rows):
                        for ci, val in enumerate(row):
                            if ci < len(table_headers):
                                table.rows[ri + 1].cells[ci].text = val
                table_rows = []
                table_headers = []
            continue

        # Image: ![alt](path)
        im = re.match(r"!\[.*?\]\((.+?)\)", line)
        if im and os.path.exists(im.group(1)):
            doc.add_picture(im.group(1), width=Inches(5))
            i += 1
            continue

        # List items
        if re.match(r"^\s*[-*+]\s+", line):
            text = re.sub(r"^\s*[-*+]\s+", "", line)
            doc.add_paragraph(text, style='List Bullet')
            i += 1
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            doc.add_paragraph(text, style='List Number')
            i += 1
            continue

        # Paragraph (skip empty lines)
        if line.strip():
            p = doc.add_paragraph(line)
            # Bold **text**
            for run in p.runs:
                if "**" in run.text:
                    run.text = run.text.replace("**", "")
                    run.bold = True
        else:
            if table_rows:
                pass  # already handled

        i += 1

    output = args.output or re.sub(r"\.md$", ".docx", args.input, flags=re.I)
    doc.save(output)
    print(json.dumps({"status": "ok", "file": output}))


# ── Argument parser ────────────────────────────────────────────────────────────
def main():
    parser = argparse.ArgumentParser(description="Document Tools")
    sub = parser.add_subparsers(dest="command", required=True)

    # docx
    d = sub.add_parser("docx")
    d_sub = d.add_subparsers(dest="action", required=True)

    dp = d_sub.add_parser("create")
    dp.add_argument("output")
    dp.add_argument("--title")
    dp.add_argument("--heading")
    dp.add_argument("--body")
    dp.add_argument("--json", help="JSON file with document structure")

    dp = d_sub.add_parser("edit")
    dp.add_argument("input")
    dp.add_argument("--replace", action="append", help="old=new pairs")
    dp.add_argument("--append")

    dp = d_sub.add_parser("to-md")
    dp.add_argument("input")
    dp.add_argument("--output")

    # xlsx
    x = sub.add_parser("xlsx")
    x_sub = x.add_subparsers(dest="action", required=True)

    xp = x_sub.add_parser("create")
    xp.add_argument("output")
    xp.add_argument("--json", help="JSON array of rows")
    xp.add_argument("--csv", help="Import from CSV")
    xp.add_argument("--rows", nargs="*", help="Row data")

    xp = x_sub.add_parser("edit")
    xp.add_argument("input")
    xp.add_argument("--set", action="append", help="Cell=value")
    xp.add_argument("--add-row")

    xp = x_sub.add_parser("convert")
    xp.add_argument("input")
    xp.add_argument("--output")

    # pdf
    p = sub.add_parser("pdf")
    p_sub = p.add_subparsers(dest="action", required=True)

    pp = p_sub.add_parser("read")
    pp.add_argument("input")
    pp.add_argument("--output")

    pp = p_sub.add_parser("merge")
    pp.add_argument("--inputs", nargs="+", required=True)
    pp.add_argument("-o", "--output")

    pp = p_sub.add_parser("split")
    pp.add_argument("input")
    pp.add_argument("--pages", help="e.g. 1-3,5")
    pp.add_argument("-o", "--output")

    pp = p_sub.add_parser("watermark")
    pp.add_argument("input")
    pp.add_argument("-w", "--watermark-text", default="DRAFT")
    pp.add_argument("-o", "--output")

    pp = p_sub.add_parser("encrypt")
    pp.add_argument("input")
    pp.add_argument("-p", "--password", required=True)
    pp.add_argument("-o", "--output")

    # markdown
    m = sub.add_parser("md")
    m_sub = m.add_subparsers(dest="action", required=True)

    mp = m_sub.add_parser("to-docx")
    mp.add_argument("input")
    mp.add_argument("-o", "--output")

    args = parser.parse_args()

    # Dispatch
    if args.command == "docx":
        if args.action == "create":
            docx_create(args)
        elif args.action == "edit":
            docx_edit(args)
        elif args.action == "to-md":
            docx_to_md(args)
    elif args.command == "xlsx":
        if args.action == "create":
            xlsx_create(args)
        elif args.action == "edit":
            xlsx_edit(args)
        elif args.action == "convert":
            xlsx_convert(args)
    elif args.command == "pdf":
        if args.action == "read":
            pdf_read(args)
        elif args.action == "merge":
            pdf_merge(args)
        elif args.action == "split":
            pdf_split(args)
        elif args.action == "watermark":
            pdf_watermark(args)
        elif args.action == "encrypt":
            pdf_encrypt(args)
    elif args.command == "md":
        if args.action == "to-docx":
            md_to_docx(args)


if __name__ == "__main__":
    main()
